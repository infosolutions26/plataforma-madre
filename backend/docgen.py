"""
Genera el entregable de contabilidad como Google Doc, RELLENANDO la plantilla de
membrete del usuario en su lugar (no recreándola).

Diseño (WYSIWYG con la vista de armado):
- La plantilla define la distribución (encabezado/pie + tablas de ejemplo en
  posiciones fijas + placeholders de texto).
- El frontend manda los BLOQUES que quedaron en la vista de armado (con las
  ediciones del encargado: renglones/tablas quitados). Cada bloque de tabla trae
  un `template_key` que dice a qué tabla de la plantilla corresponde (o null si
  no está en la plantilla y hay que ANEXARLA al final).
- Aquí: se rellenan las celdas de cada tabla de la plantilla desde los datos
  calculados, pero SOLO los renglones que siguen presentes en el bloque (los que
  el encargado quitó, se quitan del Doc); las tablas cuyo bloque ya no vino, se
  eliminan; los bloques sin template_key se anexan como tablas nuevas; las
  gráficas de ejemplo se reemplazan por las del navegador.

Solo usa el scope de Drive (exportar plantilla a docx + subir convertido a Doc).
"""
from __future__ import annotations

import base64
import copy
import io
import re
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from googleapiclient.http import MediaIoBaseUpload

import drive

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GDOC_MIME = "application/vnd.google-apps.document"

MESES = {"ENERO": 1, "FEBRERO": 2, "MARZO": 3, "ABRIL": 4, "MAYO": 5, "JUNIO": 6,
         "JULIO": 7, "AGOSTO": 8, "SEPTIEMBRE": 9, "OCTUBRE": 10, "NOVIEMBRE": 11, "DICIEMBRE": 12}


def _money(v) -> str:
    return "$" + f"{float(v or 0):,.0f}"


def _pct(v) -> str:
    return f"{float(v or 0):.2f}%"


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().upper())


def _match_key(label: str) -> str:
    """Normaliza la etiqueta de un renglón para emparejar bloque↔plantilla: si es
    un mes (ej. 'Enero 2026' o 'ENERO'), lo reduce al nombre del mes; si no, el
    texto en mayúsculas (categoría o nombre de colaborador)."""
    u = _norm(label)
    for mes in MESES:
        if mes in u:
            return mes
    return u


def _set_cell(cell, text: str):
    text = str(text)
    paras = cell.paragraphs
    if paras and paras[0].runs:
        paras[0].runs[0].text = text
        for r in paras[0].runs[1:]:
            r.text = ""
    elif paras:
        paras[0].add_run(text)
    else:
        cell.add_paragraph(text)


def _cols(tabla) -> dict:
    return {_norm(c.text): i for i, c in enumerate(tabla.rows[0].cells)}


def _quitar_tabla(tabla):
    tabla._tbl.getparent().remove(tabla._tbl)


def _quitar_fila(tabla, row):
    tabla._tbl.remove(row._tr)


# ---- rellenado de cada tipo de tabla (solo los renglones incluidos) ----

def _fill_mensual(t, datos, incluidos):
    c = _cols(t)
    ci_dep = c.get("DEPÓSITOS", c.get("DEPOSITOS")); ci_gas = c.get("GASTOS"); ci_sal = c.get("SALDO FINAL")
    monthly = datos["monthly"]
    for row in list(t.rows[1:]):
        etq = _norm(row.cells[0].text)
        if etq in MESES:
            if etq not in incluidos:
                _quitar_fila(t, row); continue
            d = monthly.get(MESES[etq])
            if ci_dep is not None: _set_cell(row.cells[ci_dep], _money(d["dep"]) if d else "$0")
            if ci_gas is not None: _set_cell(row.cells[ci_gas], _money(d["gasto"]) if d else "$0")
            if ci_sal is not None: _set_cell(row.cells[ci_sal], _money(d["saldo"]) if d else "$0")
        elif etq == "TOTAL":
            if ci_dep is not None: _set_cell(row.cells[ci_dep], _money(datos["total_dep"]))
            if ci_gas is not None: _set_cell(row.cells[ci_gas], _money(datos["total_gasto"]))


def _fill_cat_mes(t, datos, incluidos):
    c = _cols(t)
    mes_cols = {MESES[k]: v for k, v in c.items() if k in MESES}
    cat_mes = {_norm(k): v for k, v in datos["cat_mes"].items()}
    for row in list(t.rows[1:]):
        cat = _norm(row.cells[0].text)
        if cat not in incluidos:
            _quitar_fila(t, row); continue
        porm = cat_mes.get(cat, {})
        for mnum, ci in mes_cols.items():
            _set_cell(row.cells[ci], _money(porm.get(mnum, 0)))


def _fill_cat_totales(t, datos, incluidos):
    c = _cols(t)
    ci_tot = c.get("TOTAL"); ci_pct = c.get("PORCENTAJES", c.get("PORCENTAJE"))
    cat_tot = {_norm(k): v for k, v in datos["cat_tot"].items()}
    total = datos["total_gastos"] or 1.0
    for row in list(t.rows[1:]):
        cat = _norm(row.cells[0].text)
        if cat not in incluidos:
            _quitar_fila(t, row); continue
        val = cat_tot.get(cat, 0.0)
        if ci_tot is not None: _set_cell(row.cells[ci_tot], _money(val))
        if ci_pct is not None: _set_cell(row.cells[ci_pct], _pct(val / total * 100))


def _fill_dist_mes(t, datos, incluidos):
    c = _cols(t)
    ci_pct = c.get("PORCENTAJES", c.get("PORCENTAJE"))
    dist = datos["dist_mes"]
    for row in list(t.rows[1:]):
        etq = _norm(row.cells[0].text)
        if etq in MESES:
            if etq not in incluidos:
                _quitar_fila(t, row); continue
            if ci_pct is not None: _set_cell(row.cells[ci_pct], _pct(dist.get(MESES[etq], 0)))


def _fill_colaboradores(t, datos, incluidos):
    colab = [c for c in datos["colaboradores"] if _match_key(c["nombre"]) in incluidos]
    filas = t.rows[1:]
    if filas and len(colab) > len(filas):
        molde = filas[-1]._tr
        for _ in range(len(colab) - len(filas)):
            t._tbl.append(copy.deepcopy(molde))
    filas = t.rows[1:]
    for idx, row in enumerate(list(filas)):
        if idx < len(colab):
            _set_cell(row.cells[0], colab[idx]["nombre"])
            _set_cell(row.cells[1], _money(colab[idx]["total"]))
        else:
            _quitar_fila(t, row)


def _tabla_key(heads: set) -> Optional[str]:
    if ({"DEPÓSITOS", "SALDO FINAL"} & heads) or ({"DEPOSITOS", "SALDO FINAL"} & heads):
        return "mensual"
    if "GASTOS" in heads and (heads & set(MESES.keys())):
        return "cat_mes"
    if "GASTOS" in heads and ("PORCENTAJES" in heads or "PORCENTAJE" in heads) and "TOTAL" in heads:
        return "cat_totales"
    if "MES" in heads and ("PORCENTAJES" in heads or "PORCENTAJE" in heads):
        return "dist_mes"
    if {"NOMBRE", "MONTO"} <= heads:
        return "colaboradores"
    return None


_FILLERS = {"mensual": _fill_mensual, "cat_mes": _fill_cat_mes, "cat_totales": _fill_cat_totales,
            "dist_mes": _fill_dist_mes, "colaboradores": _fill_colaboradores}


def _rellenar_tablas(doc, datos, secciones):
    """secciones: {template_key: set(match_keys incluidos)}. Tabla sin sección -> se quita."""
    for t in list(doc.tables):
        key = _tabla_key(set(_cols(t).keys()))
        if key is None:
            continue
        if key not in secciones:
            _quitar_tabla(t)   # el encargado la excluyó de la vista de armado
        else:
            _FILLERS[key](t, datos, secciones[key])


def _reemplazar_texto(doc, mapa: dict):
    for p in doc.paragraphs:
        for clave, valor in mapa.items():
            if clave and clave in p.text and valor:
                nuevo = p.text.replace(clave, str(valor))
                if p.runs:
                    p.runs[0].text = nuevo
                    for r in p.runs[1:]:
                        r.text = ""
                break


def _titulo(doc, texto, antes=None):
    p = antes.insert_paragraph_before() if antes is not None else doc.add_paragraph()
    r = p.add_run(texto); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x07, 0x2E, 0x7D)
    return p


def _anexar_tabla(doc, bloque, antes):
    """Inserta una tabla nueva (top-10, comparativo) antes del párrafo `antes`."""
    _titulo(doc, bloque.get("titulo", ""), antes)
    cols = bloque.get("columnas", []); filas = bloque.get("filas", [])
    t = doc.add_table(rows=1, cols=max(1, len(cols)))
    try:
        t.style = "Table Grid"
    except KeyError:
        pass
    for i, c in enumerate(cols):
        _set_cell(t.rows[0].cells[i], c)
        for par in t.rows[0].cells[i].paragraphs:
            for run in par.runs:
                run.bold = True
    for fila in filas:
        cells = t.add_row().cells
        for i, val in enumerate(fila):
            if i < len(cells):
                _set_cell(cells[i], val)
    # mover la tabla (que add_table pone al final) a la posición antes del cierre
    if antes is not None:
        antes._p.addprevious(t._tbl)


def _reemplazar_graficas(doc, graficas_png_b64):
    try:
        shapes = doc.inline_shapes
    except Exception:  # noqa: BLE001
        return
    for i, png_b64 in enumerate(graficas_png_b64):
        if i >= len(shapes):
            break
        try:
            rId = shapes[i]._inline.graphic.graphicData.pic.blipFill.blip.embed
            doc.part.related_parts[rId]._blob = base64.b64decode(png_b64)
        except Exception:  # noqa: BLE001
            continue


def rellenar_plantilla(
    plantilla_doc_id: str,
    nombre_cliente: str,
    periodo: str,
    fecha: str,
    datos: dict,
    secciones: dict,          # {template_key: [labels incluidos]}
    extras: list,             # bloques de tabla a anexar (top-10, comparativo)
    graficas_png: list,       # PNG base64 en orden
    notas: str,
    carpeta_id: Optional[str] = None,
) -> dict:
    svc = drive._get_service()
    docx_bytes = svc.files().export(fileId=plantilla_doc_id, mimeType=DOCX_MIME).execute()
    doc = Document(io.BytesIO(docx_bytes))

    _reemplazar_texto(doc, {
        "NOMBRE DE COMPAÑÍA": nombre_cliente,
        "Mes - Mes": periodo,
        "Mes y Día, 2026": fecha,
    })

    sec_norm = {k: set(_match_key(l) for l in labels) for k, labels in secciones.items()}
    _rellenar_tablas(doc, datos, sec_norm)
    _reemplazar_graficas(doc, graficas_png or [])

    # punto de inserción para lo que se anexa: antes del cierre "Quedamos atentos..."
    cierre = None
    for p in doc.paragraphs:
        if "QUEDAMOS ATENTOS" in _norm(p.text):
            cierre = p
            break
    for bloque in (extras or []):
        if cierre is not None:
            _anexar_tabla(doc, bloque, cierre)
        else:
            _titulo(doc, bloque.get("titulo", ""))
            _anexar_tabla(doc, bloque, None)

    if notas and notas.strip():
        for linea in notas.split("\n"):
            if cierre is not None:
                cierre.insert_paragraph_before(linea)
            else:
                doc.add_paragraph(linea)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    media = MediaIoBaseUpload(out, mimetype=DOCX_MIME, resumable=False)
    meta = {"name": f"Entregable Contabilidad — {nombre_cliente} ({periodo})", "mimeType": GDOC_MIME}
    if carpeta_id:
        meta["parents"] = [carpeta_id]
    f = svc.files().create(body=meta, media_body=media, fields="id, webViewLink").execute()
    try:
        svc.permissions().create(fileId=f["id"], body={"type": "anyone", "role": "writer"}, fields="id").execute()
    except Exception:  # noqa: BLE001
        pass
    return f
